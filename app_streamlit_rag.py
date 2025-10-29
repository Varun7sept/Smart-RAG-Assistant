# -------------------------------
# 🤖 Smart RAG Chat Assistant (Multi-user + MongoDB + Multi-format Upload + Chat Logging)
# Persistent + Conversational + Upload + Delete + Viewer + Memory Reset + Save Logs
# -------------------------------

import streamlit as st
from dotenv import load_dotenv
import os
from pathlib import Path
import hashlib
from datetime import datetime
from pymongo import MongoClient
import bcrypt
import pandas as pd
from langchain.schema import Document
from docx import Document as DocxDocument

# -------------------------------
# Load environment variables
# -------------------------------
load_dotenv()

# -------------------------------
# 🔌 MongoDB Connection
# -------------------------------
mongo_uri = os.getenv("MONGO_URI")
db_name = os.getenv("DB_NAME")

try:
    client = MongoClient(mongo_uri)
    db = client[db_name]
    users_collection = db["users"]
    st.sidebar.success("✅ Connected to MongoDB")
except Exception as e:
    st.sidebar.error(f"⚠️ MongoDB connection failed: {e}")
    st.stop()

# -------------------------------
# 🔐 User Authentication
# -------------------------------
st.sidebar.header("🔐 User Authentication")

auth_choice = st.sidebar.radio("Choose an option:", ["Login", "Sign Up"])

if auth_choice == "Sign Up":
    st.subheader("🆕 Create New Account")
    new_username = st.text_input("Username")
    new_email = st.text_input("Email")
    new_password = st.text_input("Password", type="password")
    confirm_password = st.text_input("Confirm Password", type="password")

    if st.button("Register"):
        if new_password != confirm_password:
            st.error("❌ Passwords do not match.")
        elif users_collection.find_one({"username": new_username}):
            st.warning("⚠️ Username already exists.")
        elif users_collection.find_one({"email": new_email}):
            st.warning("⚠️ Email already registered.")
        else:
            hashed_pw = bcrypt.hashpw(new_password.encode(), bcrypt.gensalt()).decode()
            users_collection.insert_one({
                "username": new_username,
                "email": new_email,
                "password": hashed_pw
            })
            st.success("✅ Registration successful! Please login now.")
            st.rerun()

elif auth_choice == "Login":
    st.subheader("🔑 Login to Your Account")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")

    if st.button("Login"):
        user = users_collection.find_one({"username": username})
        if not user:
            st.error("❌ User not found.")
        elif bcrypt.checkpw(password.encode(), user["password"].encode()):
            st.session_state["authenticated"] = True
            st.session_state["username"] = user["username"]
            st.session_state["email"] = user["email"]
            st.success(f"✅ Welcome back, {username}!")
            st.rerun()
        else:
            st.error("❌ Incorrect password.")

# -------------------------------
# 🚫 Prevent Unauthenticated Access
# -------------------------------
if "authenticated" not in st.session_state or not st.session_state["authenticated"]:
    st.warning("🔒 Please log in or sign up to continue.")
    st.stop()

# -------------------------------
# 🧠 Session & User Data Setup
# -------------------------------
username = st.session_state["username"]
email = st.session_state["email"]
session_id = email.split("@")[0]

st.sidebar.markdown(f"👤 Logged in as: `{username}`")

if st.sidebar.button("🚪 Logout"):
    st.session_state.clear()
    st.success("👋 Logged out successfully.")
    st.rerun()

# -------------------------------
# LangChain + RAG Imports
# -------------------------------
from langchain_groq import ChatGroq
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    CSVLoader,
    JSONLoader,
    UnstructuredHTMLLoader
)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain

# -------------------------------
# Helper — Compute MD5 Hash
# -------------------------------
def compute_md5(file_path: Path) -> str:
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()

# -------------------------------
# Helper — Load Different File Types
# -------------------------------
def load_document(file_path: Path):
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return PyPDFLoader(str(file_path)).load()
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return [Document(page_content=text, metadata={"source": str(file_path)})]
    elif ext == ".csv":
        df = pd.read_csv(file_path)
        return [Document(page_content=df.to_string(index=False), metadata={"source": str(file_path)})]
    elif ext == ".json":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return [Document(page_content=text, metadata={"source": str(file_path)})]
    elif ext == ".html":
        return UnstructuredHTMLLoader(str(file_path)).load()
    elif ext == ".docx":
        doc = DocxDocument(str(file_path))
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return [Document(page_content=text, metadata={"source": str(file_path)})]
    else:
        st.warning(f"⚠️ Unsupported file type: {ext}")
        return []

# -------------------------------
# Streamlit Setup
# -------------------------------
st.set_page_config(page_title="Smart RAG Chat Assistant 🤖", layout="wide")
st.title("🤖 Smart RAG Chat Assistant")
st.markdown("Chat with **PDF, DOCX, TXT, CSV, JSON, or HTML** files using **Groq + LangChain + MongoDB + ChromaDB + HuggingFace**.")

# -------------------------------
# Directory Setup (Per User)
# -------------------------------
base_dir = Path(__file__).parent
data_dir = base_dir / "data"
chroma_dir = base_dir / "chroma_db"
logs_dir = base_dir / "logs"

for d in [data_dir, chroma_dir, logs_dir]:
    d.mkdir(exist_ok=True)

user_data_dir = data_dir / session_id
user_chroma_dir = chroma_dir / session_id
user_logs_dir = logs_dir / session_id

for d in [user_data_dir, user_chroma_dir, user_logs_dir]:
    d.mkdir(exist_ok=True)

# -------------------------------
# Chat Logger
# -------------------------------
def log_chat(user, question, answer):
    log_file = user_logs_dir / f"chat_{datetime.now().strftime('%Y-%m-%d')}.txt"
    with open(log_file, "a", encoding="utf-8") as f:
        f.write(f"[{datetime.now().strftime('%H:%M:%S')}] USER: {question}\n")
        f.write(f"[{datetime.now().strftime('%H:%M:%S')}] ASSISTANT: {answer}\n\n")

# -------------------------------
# Initialize Embeddings + DB
# -------------------------------
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
db = Chroma(persist_directory=str(user_chroma_dir), embedding_function=embeddings)

# -------------------------------
# Sidebar Tabs
# -------------------------------
tabs = st.sidebar.radio(
    "📚 Navigation",
    ["📂 Manage Files", "💬 Chat Interface", "🧠 Knowledge Base Viewer", "📜 View Chat Logs"],
    index=1
)

# -------------------------------
# 📂 Manage Files
# -------------------------------
if tabs == "📂 Manage Files":
    st.sidebar.header("📂 Upload or Delete Documents")

    uploaded_files = st.sidebar.file_uploader(
        "Upload one or more files",
        type=["pdf", "txt", "csv", "json", "html", "docx"],
        accept_multiple_files=True
    )

    if uploaded_files:
        for uploaded_file in uploaded_files:
            file_path = user_data_dir / uploaded_file.name
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            st.sidebar.success(f"✅ Saved: {uploaded_file.name}")

    existing_files = list(user_data_dir.glob("*"))
    if existing_files:
        st.sidebar.subheader("🗂 Current Files:")
        for file in existing_files:
            col1, col2 = st.sidebar.columns([4, 1])
            col1.write(f"- {file.name}")
            if col2.button("❌", key=str(file)):
                db.delete(where={"source": str(file)})
                os.remove(file)
                db.persist()
                st.sidebar.warning(f"🗑 Deleted {file.name}")
                st.rerun()
    else:
        st.sidebar.info("No files found. Upload some to get started!")

# -------------------------------
# 💾 Process New / Updated Files
# -------------------------------
existing_meta = db.get(include=["metadatas"]).get("metadatas", [])
existing_hashes = {m["source"]: m["hash"] for m in existing_meta if "source" in m and "hash" in m}

files = list(user_data_dir.glob("*"))
to_process, to_delete = [], []

for file in files:
    file_hash = compute_md5(file)
    if str(file) not in existing_hashes or existing_hashes[str(file)] != file_hash:
        to_process.append((file, file_hash))
        if str(file) in existing_hashes:
            to_delete.append(str(file))

if to_delete:
    for f in to_delete:
        db.delete(where={"source": f})

if to_process:
    st.sidebar.info(f"📥 Processing {len(to_process)} new/updated files...")
    all_docs = []
    for file, file_hash in to_process:
        docs = load_document(file)
        for d in docs:
            d.metadata = {"source": str(file), "hash": file_hash}
        all_docs.extend(docs)

    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100)
    chunks = splitter.split_documents(all_docs)
    db.add_documents(chunks)
    db.persist()
    st.sidebar.success("💾 Your private ChromaDB updated successfully!")

retriever = db.as_retriever(search_kwargs={"k": 3})

# -------------------------------
# 🧠 Knowledge Base Viewer
# -------------------------------
if tabs == "🧠 Knowledge Base Viewer":
    st.header("🧠 Knowledge Base Viewer")
    meta_data = db.get(include=["metadatas", "documents"])
    if not meta_data["documents"]:
        st.info("📭 No documents found. Upload files and process them first.")
    else:
        for i, (doc, meta) in enumerate(zip(meta_data["documents"], meta_data["metadatas"]), start=1):
            st.markdown(f"### 🧩 Chunk {i}")
            st.markdown(f"**Source:** {Path(meta['source']).name}")
            st.markdown(f"**Excerpt:** {doc[:300]}...")
            st.divider()

# -------------------------------
# 💬 Chat Interface
# -------------------------------
if tabs == "💬 Chat Interface":
    st.header("💬 Chat with Your Documents")

    # ✅ Updated models (deprecated ones removed)
    model_choice = st.selectbox(
        "Select Model:",
        ["llama-3.1-8b-instant", "llama-3.1-70b-versatile", "mixtral-8x22b", "gemma2-9b-it"],
        index=0
    )

    llm = ChatGroq(model=model_choice, groq_api_key=os.getenv("GROQ_API_KEY"))
    memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)
    qa_chain = ConversationalRetrievalChain.from_llm(llm=llm, retriever=retriever, memory=memory)

    if st.button("🧹 Clear Chat Memory"):
        st.session_state.history = []
        memory.clear()
        st.info("🧽 Chat memory cleared!")

    if "history" not in st.session_state:
        st.session_state.history = []

    user_input = st.chat_input("Ask a question about your files...")
    if user_input:
        try:
            result = qa_chain.invoke({"question": user_input})
            answer = result["answer"]
            st.session_state.history.append(("user", user_input))
            st.session_state.history.append(("assistant", answer))
            log_chat(username, user_input, answer)
        except Exception as e:
            st.error(f"⚠️ Groq API Error: {e}")
            st.info("Try selecting another available model.")

    for role, message in st.session_state.history:
        st.chat_message(role).markdown(f"**{role.capitalize()}:** {message}")

# -------------------------------
# 📜 View Chat Logs
# -------------------------------
if tabs == "📜 View Chat Logs":
    st.header("📜 View Your Chat Logs")
    log_files = sorted(user_logs_dir.glob("*.txt"))
    if not log_files:
        st.info("No logs found yet. Start chatting first!")
    else:
        selected_log = st.selectbox("Select log file:", [f.name for f in log_files])
        if selected_log:
            with open(user_logs_dir / selected_log, "r", encoding="utf-8") as f:
                st.text(f.read())
