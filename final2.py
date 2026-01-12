# -------------------------------
# 🤖 Smart RAG Chat Assistant (Multi-user + MongoDB + Multi-format Upload + Chat Logging)
# Persistent + Conversational + Upload + Delete + Viewer + Memory Reset + Save Logs
# + RAG Evaluation Metrics + Inline Citations + MMR Retrieval (Enhanced)
# -------------------------------

import streamlit as st
from dotenv import load_dotenv
import os
from pathlib import Path
import hashlib
from datetime import datetime
from pymongo import MongoClient
import bcrypt
import pandas as pd
import numpy as np
from langchain.schema import Document
from docx import Document as DocxDocument

# -------------------------------
# Load environment variables
# -------------------------------
load_dotenv()

# -------------------------------
# 🔌 MongoDB Connection
# -------------------------------
mongo_uri = os.getenv("MONGO_URI")
db_name = os.getenv("DB_NAME")

try:
    client = MongoClient(mongo_uri)
    db = client[db_name]
    users_collection = db["users"]
    st.sidebar.success("✅ Connected to MongoDB")
except Exception as e:
    st.sidebar.error(f"⚠️ MongoDB connection failed: {e}")
    st.stop()

# -------------------------------
# 🔐 User Authentication
# -------------------------------
st.sidebar.header("🔐 User Authentication")

auth_choice = st.sidebar.radio("Choose an option:", ["Login", "Sign Up"])

if auth_choice == "Sign Up":
    st.subheader("🆕 Create New Account")
    new_username = st.text_input("Username")
    new_email = st.text_input("Email")
    new_password = st.text_input("Password", type="password")
    confirm_password = st.text_input("Confirm Password", type="password")

    if st.button("Register"):
        if new_password != confirm_password:
            st.error("❌ Passwords do not match.")
        elif users_collection.find_one({"username": new_username}):
            st.warning("⚠️ Username already exists.")
        elif users_collection.find_one({"email": new_email}):
            st.warning("⚠️ Email already registered.")
        else:
            hashed_pw = bcrypt.hashpw(new_password.encode(), bcrypt.gensalt()).decode()
            users_collection.insert_one({
                "username": new_username,
                "email": new_email,
                "password": hashed_pw
            })
            st.success("✅ Registration successful! Please login now.")
            st.rerun()

elif auth_choice == "Login":
    st.subheader("🔑 Login to Your Account")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")

    if st.button("Login"):
        user = users_collection.find_one({"username": username})
        if not user:
            st.error("❌ User not found.")
        elif bcrypt.checkpw(password.encode(), user["password"].encode()):
            st.session_state["authenticated"] = True
            st.session_state["username"] = user["username"]
            st.session_state["email"] = user["email"]
            st.success(f"✅ Welcome back, {username}!")
            st.rerun()
        else:
            st.error("❌ Incorrect password.")

# -------------------------------
# 🚫 Prevent Unauthenticated Access
# -------------------------------
if "authenticated" not in st.session_state or not st.session_state["authenticated"]:
    st.warning("🔒 Please log in or sign up to continue.")
    st.stop()

# -------------------------------
# 🧠 Session & User Data Setup
# -------------------------------
username = st.session_state["username"]
email = st.session_state["email"]
session_id = email.split("@")[0]

st.sidebar.markdown(f"👤 Logged in as: `{username}`")

if st.sidebar.button("🚪 Logout"):
    st.session_state.clear()
    st.success("👋 Logged out successfully.")
    st.rerun()


# -------------------------------
# LangChain + RAG Imports
# -------------------------------
from langchain_groq import ChatGroq
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.prompts import PromptTemplate
from langchain_community.document_loaders import (
    PyPDFLoader,
    UnstructuredHTMLLoader
)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain

# -------------------------------
# Helper — Compute MD5 Hash
# -------------------------------
def compute_md5(file_path: Path) -> str:
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()

# -------------------------------
# Helper — Load Different File Types
# -------------------------------
def load_document(file_path: Path):
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return PyPDFLoader(str(file_path)).load()
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return [Document(page_content=text, metadata={"source": str(file_path)})]
    elif ext == ".csv":
        df = pd.read_csv(file_path)
        return [Document(page_content=df.to_string(index=False), metadata={"source": str(file_path)})]
    elif ext == ".json":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return [Document(page_content=text, metadata={"source": str(file_path)})]
    elif ext == ".html":
        return UnstructuredHTMLLoader(str(file_path)).load()
    elif ext == ".docx":
        doc = DocxDocument(str(file_path))
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return [Document(page_content=text, metadata={"source": str(file_path)})]
    else:
        st.warning(f"⚠️ Unsupported file type: {ext}")
        return []

# -------------------------------
# 🧭 Query Intent Detection (NEW)
# -------------------------------
def detect_intent(q: str) -> str:
    ql = q.lower()

    if any(k in ql for k in ["summary", "summarize", "overall", "entire", "document", "resume", "profile"]):
        return "summary"

    if any(k in ql for k in ["compare", "difference", "vs", "versus"]):
        return "compare"

    if any(k in ql for k in ["explain", "in simple", "simple words", "what is"]):
        return "explain"

    return "fact"  # default intent

# -------------------------------
# 🔬 Improved RAG Evaluation Utilities
# -------------------------------
def cosine_sim(a: np.ndarray, b: np.ndarray) -> float:
    if np.linalg.norm(a) == 0 or np.linalg.norm(b) == 0:
        return 0.0
    return float(np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b)))

def compute_rag_metrics(question: str, answer: str, contexts: list[str], embedder):
    """
    Returns normalized 0–100 scores:
      - Answer Relevance
      - Context Relevance (adaptive: broad vs focused)
      - Faithfulness
      - Overall Score
    """
    q_emb = np.array(embedder.embed_query(question))
    a_emb = np.array(embedder.embed_query(answer))

    # Answer relevance (Q ↔ A)
    ans_rel = cosine_sim(q_emb, a_emb)

    # Decide if the query is broad (document-level)
    broad_query = len(question.split()) <= 6 and any(
        k in question.lower()
        for k in ["document", "resume", "profile", "overall", "summary", "entire"]
    )

    if broad_query and contexts:
        # Treat all context as one combined document
        combined_ctx = " ".join(contexts)
        c_emb = np.array(embedder.embed_query(combined_ctx))
        ctx_rel = cosine_sim(q_emb, c_emb)
    else:
        # Normal per-chunk evaluation (best matching chunk)
        ctx_sims = []
        for ctx in contexts:
            c_emb = np.array(embedder.embed_query(ctx))
            ctx_sims.append(cosine_sim(q_emb, c_emb))
        ctx_rel = float(np.max(ctx_sims)) if ctx_sims else 0.0

    # Faithfulness (A ↔ combined context)
    if contexts:
        combined_ctx = " ".join(contexts)
        c_all_emb = np.array(embedder.embed_query(combined_ctx))
        faithful = cosine_sim(a_emb, c_all_emb)
    else:
        faithful = 0.0

    ans_rel_n = round(ans_rel * 100, 2)
    ctx_rel_n = round(ctx_rel * 100, 2)
    faithful_n = round(faithful * 100, 2)

    overall = round((0.4 * ans_rel_n) + (0.3 * ctx_rel_n) + (0.3 * faithful_n), 2)

    return {
        "answer_relevance": ans_rel_n,
        "context_relevance": ctx_rel_n,
        "faithfulness": faithful_n,
        "overall": overall
    }


# -------------------------------
# Streamlit Setup
# -------------------------------
st.set_page_config(page_title="Smart RAG Chat Assistant 🤖", layout="wide")
st.title("🤖 Smart RAG Chat Assistant")
st.markdown(
    "Chat with **PDF, DOCX, TXT, CSV, JSON, or HTML** files using "
    "**Groq + LangChain + MongoDB + ChromaDB + HuggingFace**."
)

# -------------------------------
# Directory Setup (Per User)
# -------------------------------
base_dir = Path(__file__).parent
data_dir = base_dir / "data"
chroma_dir = base_dir / "chroma_db"
logs_dir = base_dir / "logs"

for d in [data_dir, chroma_dir, logs_dir]:
    d.mkdir(exist_ok=True)

user_data_dir = data_dir / session_id
user_chroma_dir = chroma_dir / session_id
user_logs_dir = logs_dir / session_id

for d in [user_data_dir, user_chroma_dir, user_logs_dir]:
    d.mkdir(exist_ok=True)

# -------------------------------
# Chat Logger
# -------------------------------
def log_chat(user, question, answer):
    log_file = user_logs_dir / f"chat_{datetime.now().strftime('%Y-%m-%d')}.txt"
    with open(log_file, "a", encoding="utf-8") as f:
        f.write(f"[{datetime.now().strftime('%H:%M:%S')}] USER: {question}\n")
        f.write(f"[{datetime.now().strftime('%H:%M:%S')}] ASSISTANT: {answer}\n\n")

# -------------------------------
# Initialize Embeddings + Vector DB
# -------------------------------
embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)

db = Chroma(
    persist_directory=str(user_chroma_dir),
    embedding_function=embeddings
)

# -------------------------------
# Sidebar Tabs
# -------------------------------
tabs = st.sidebar.radio(
    "📚 Navigation",
    [
        "📂 Manage Files",
        "💬 Chat Interface",
        "🧠 Knowledge Base Viewer",
        "📊 RAG Evaluation Metrics",
        "📜 View Chat Logs"
    ],
    index=1
)

# Initialize state containers
if "history" not in st.session_state:
    st.session_state["history"] = []

if "rag_metrics" not in st.session_state:
    st.session_state["rag_metrics"] = []


# -------------------------------
# 📂 Manage Files
# -------------------------------
if tabs == "📂 Manage Files":
    st.sidebar.header("📂 Upload or Delete Documents")

    uploaded_files = st.sidebar.file_uploader(
        "Upload one or more files",
        type=["pdf", "txt", "csv", "json", "html", "docx"],
        accept_multiple_files=True
    )

    if uploaded_files:
        for uploaded_file in uploaded_files:
            file_path = user_data_dir / uploaded_file.name
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            st.sidebar.success(f"✅ Saved: {uploaded_file.name}")

    existing_files = list(user_data_dir.glob("*"))
    if existing_files:
        st.sidebar.subheader("🗂 Current Files:")
        for file in existing_files:
            col1, col2 = st.sidebar.columns([4, 1])
            col1.write(f"- {file.name}")
            if col2.button("❌", key=str(file)):
                db.delete(where={"source": str(file)})
                os.remove(file)
                db.persist()
                st.sidebar.warning(f"🗑 Deleted {file.name}")
                st.rerun()
    else:
        st.sidebar.info("No files found. Upload some to get started!")

# -------------------------------
# 💾 Process New / Updated Files
# -------------------------------
existing_meta = db.get(include=["metadatas"]).get("metadatas", [])
existing_hashes = {
    m["source"]: m["hash"]
    for m in existing_meta
    if "source" in m and "hash" in m
}

files = list(user_data_dir.glob("*"))
to_process, to_delete = [], []

for file in files:
    file_hash = compute_md5(file)
    if str(file) not in existing_hashes or existing_hashes[str(file)] != file_hash:
        to_process.append((file, file_hash))
        if str(file) in existing_hashes:
            to_delete.append(str(file))

if to_delete:
    for f in to_delete:
        db.delete(where={"source": f})

if to_process:
    st.sidebar.info(f"📥 Processing {len(to_process)} new/updated files...")
    all_docs = []
    for file, file_hash in to_process:
        docs = load_document(file)
        for d in docs:
            d.metadata = {"source": str(file), "hash": file_hash}
        all_docs.extend(docs)

    # 🔧 Improved chunking for better RAG quality
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=150
    )

    chunks = splitter.split_documents(all_docs)
    db.add_documents(chunks)
    db.persist()
    st.sidebar.success("💾 Your private ChromaDB updated successfully!")

# -------------------------------
# 🔍 Retriever (MMR for diversity + higher recall)
# -------------------------------
retriever = db.as_retriever(
    search_type="mmr",
    search_kwargs={"k": 10, "fetch_k": 20}
)


# -------------------------------
# 🧠 Knowledge Base Viewer
# -------------------------------
if tabs == "🧠 Knowledge Base Viewer":
    st.header("🧠 Knowledge Base Viewer")
    meta_data = db.get(include=["metadatas", "documents"])

    if not meta_data["documents"]:
        st.info("📭 No documents found. Upload files and process them first.")
    else:
        for i, (doc, meta) in enumerate(
            zip(meta_data["documents"], meta_data["metadatas"]), start=1
        ):
            st.markdown(f"### 🧩 Chunk {i}")
            st.markdown(f"**Source:** {Path(meta['source']).name}")
            st.markdown(f"**Excerpt:** {doc[:400]}...")
            st.divider()


# -------------------------------
# 💬 Chat Interface
# -------------------------------
if tabs == "💬 Chat Interface":
    st.header("💬 Chat with Your Documents")

    ALL_MODELS = [
        "llama-3.1-8b-instant",
        "llama-3.3-70b-versatile",
        "openai/gpt-oss-20b",
        "openai/gpt-oss-120b",
        "qwen/qwen3-32b"
    ]

    compare_all = st.checkbox("🧪 Enable Multi-Model Evaluation (Leaderboard Mode)")

    # -------- Prompt Templates by Intent --------
    PROMPTS = {
        "summary": """
You are an academic assistant. Generate a comprehensive summary of the document.

Guidelines:
- Cover all major sections.
- Be structured and slightly verbose.
- Use headings or numbered points.
- Do not omit important details.

Context:
{context}

Question:
{question}

Answer:
""",
        "compare": """
You are an academic assistant. Provide a structured comparison.

Guidelines:
- Use tables or numbered comparisons if applicable.
- Clearly distinguish differences and similarities.
- Be precise and analytical.

Context:
{context}

Question:
{question}

Answer:
""",
        "explain": """
You are a helpful tutor. Explain in simple and easy words.

Guidelines:
- Avoid complex terminology.
- Use simple language and short sentences.
- Give intuitive explanations.

Context:
{context}

Question:
{question}

Answer:
""",
        "fact": """
You are an academic assistant. Answer accurately and concisely.

Guidelines:
- Stick strictly to the given documents.
- Be clear and factual.
- Do not hallucinate.

Context:
{context}

Question:
{question}

Answer:
"""
    }

    if not compare_all:
        # -------- Normal Mode (Single Model) --------
        model_choice = st.selectbox("Select Model:", ALL_MODELS)

        llm = ChatGroq(model=model_choice, groq_api_key=os.getenv("Groq_API_KEY"))

        memory = ConversationBufferMemory(
            memory_key="chat_history",
            return_messages=True,
            output_key="answer"
        )

        if st.button("🧹 Clear Chat Memory"):
            st.session_state.history = []
            st.session_state["rag_metrics"] = []
            memory.clear()
            st.info("🧽 Chat memory & metrics cleared!")

        user_input = st.chat_input("Ask a question about your files...")

        if user_input:
            intent = detect_intent(user_input)
            st.caption(f"🧭 Detected Query Type: **{intent.upper()}**")

            qa_prompt = PromptTemplate(
                input_variables=["context", "question"],
                template=PROMPTS[intent]
            )

            qa_chain = ConversationalRetrievalChain.from_llm(
                llm=llm,
                retriever=retriever,
                memory=memory,
                return_source_documents=True,
                combine_docs_chain_kwargs={"prompt": qa_prompt}
            )

            result = qa_chain.invoke({"question": user_input})

            answer = result["answer"]
            source_docs = result.get("source_documents", [])

            cited_answer = answer + "\n\n**Sources:**\n"
            contexts = []

            for i, d in enumerate(source_docs, start=1):
                src_name = Path(d.metadata.get("source", "Unknown")).name
                cited_answer += f"[{i}] {src_name}\n"
                contexts.append(d.page_content)

            metrics = compute_rag_metrics(
                question=user_input,
                answer=answer,
                contexts=contexts,
                embedder=embeddings
            )

            st.session_state.history.append(("user", user_input))
            st.session_state.history.append(("assistant", cited_answer))

            st.session_state["rag_metrics"].append({
                "mode": "single",
                "model": model_choice,
                "intent": intent,
                "question": user_input,
                "answer": cited_answer,
                "contexts": contexts,
                "metrics": metrics,
                "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            })

            log_chat(username, user_input, cited_answer)

        for role, message in st.session_state.history:
            st.chat_message(role).markdown(f"**{role.capitalize()}:** {message}")

    else:
        # -------- Multi-Model Leaderboard Mode --------
        st.info("🧪 Leaderboard Mode: The same question will be evaluated on all models.")

        user_input = st.chat_input("Ask a question to benchmark all models...")

        if user_input:
            intent = detect_intent(user_input)
            st.caption(f"🧭 Detected Query Type: **{intent.upper()}**")

            leaderboard = []

            qa_prompt = PromptTemplate(
                input_variables=["context", "question"],
                template=PROMPTS[intent]
            )

            for model_name in ALL_MODELS:
                llm = ChatGroq(model=model_name, groq_api_key=os.getenv("Groq_API_KEY"))

                memory = ConversationBufferMemory(
                    memory_key="chat_history",
                    return_messages=True,
                    output_key="answer"
                )

                qa_chain = ConversationalRetrievalChain.from_llm(
                    llm=llm,
                    retriever=retriever,
                    memory=memory,
                    return_source_documents=True,
                    combine_docs_chain_kwargs={"prompt": qa_prompt}
                )

                result = qa_chain.invoke({"question": user_input})

                answer = result["answer"]
                source_docs = result.get("source_documents", [])

                cited_answer = answer + "\n\n**Sources:**\n"
                contexts = []

                for i, d in enumerate(source_docs, start=1):
                    src_name = Path(d.metadata.get("source", "Unknown")).name
                    cited_answer += f"[{i}] {src_name}\n"
                    contexts.append(d.page_content)

                metrics = compute_rag_metrics(
                    question=user_input,
                    answer=answer,
                    contexts=contexts,
                    embedder=embeddings
                )

                leaderboard.append({
                    "model": model_name,
                    "intent": intent,
                    "question": user_input,
                    "answer": cited_answer,
                    "contexts": contexts,
                    "metrics": metrics
                })

            leaderboard.sort(key=lambda x: x["metrics"]["overall"], reverse=True)

            st.session_state["rag_metrics"].append({
                "mode": "multi",
                "intent": intent,
                "question": user_input,
                "results": leaderboard,
                "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            })

            st.success("Leaderboard generated! Open 📊 RAG Evaluation Metrics to view results.")



# -------------------------------
# 📊 RAG Evaluation Metrics
# -------------------------------
if tabs == "📊 RAG Evaluation Metrics":
    st.header("📊 RAG Evaluation Metrics")

    metrics_data = st.session_state.get("rag_metrics", [])

    if not metrics_data:
        st.info("No evaluation data yet. Ask some questions first.")
    else:
        entries = [
            f"{i+1}. [{m.get('mode','single').upper()}] {m.get('question','')[:70]}..."
            for i, m in enumerate(metrics_data)
        ]

        idx = st.selectbox(
            "Select an evaluation entry:",
            range(len(entries)),
            format_func=lambda i: entries[i]
        )

        selected = metrics_data[idx]

        # ------------------ Build Report Text ------------------
        report_lines = []
        report_lines.append("RAG EVALUATION REPORT")
        report_lines.append("=" * 60)
        report_lines.append(f"Time: {selected.get('time', '')}")
        report_lines.append("")
        report_lines.append("Question:")
        report_lines.append(selected.get("question", ""))
        report_lines.append("")

        if selected["mode"] == "single":
            m = selected["metrics"]

            st.subheader("📝 Question")
            st.write(selected["question"])

            st.subheader("🤖 Answer")
            st.write(selected["answer"])

            st.subheader("📈 RAG Quality Scores")
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Answer Relevance", m["answer_relevance"])
            c2.metric("Context Relevance", m["context_relevance"])
            c3.metric("Faithfulness", m["faithfulness"])
            c4.metric("Overall", m["overall"])

            report_lines.append("Mode: SINGLE MODEL")
            report_lines.append("")
            report_lines.append("Scores:")
            report_lines.append(f"  Answer Relevance : {m['answer_relevance']}")
            report_lines.append(f"  Context Relevance: {m['context_relevance']}")
            report_lines.append(f"  Faithfulness     : {m['faithfulness']}")
            report_lines.append(f"  Overall Score    : {m['overall']}")
            report_lines.append("")
            report_lines.append("Answer:")
            report_lines.append(selected["answer"])

        else:
            st.subheader("📝 Question")
            st.write(selected["question"])

            st.subheader("🏆 Model Leaderboard")

            table = []
            for r in selected["results"]:
                table.append({
                    "Model": r["model"],
                    "Answer Relevance": r["metrics"]["answer_relevance"],
                    "Context Relevance": r["metrics"]["context_relevance"],
                    "Faithfulness": r["metrics"]["faithfulness"],
                    "Overall": r["metrics"]["overall"]
                })

            df = pd.DataFrame(table).sort_values("Overall", ascending=False)
            st.dataframe(df, use_container_width=True)

            chosen = st.selectbox(
                "Inspect model:",
                range(len(selected["results"])),
                format_func=lambda i: selected["results"][i]["model"]
            )

            r = selected["results"][chosen]

            st.subheader(f"🤖 Answer – {r['model']}")
            st.write(r["answer"])

            report_lines.append("Mode: MULTI-MODEL LEADERBOARD")
            report_lines.append("")
            report_lines.append("Model Scores:")
            for row in df.to_dict(orient="records"):
                report_lines.append(
                    f"{row['Model']} | AR: {row['Answer Relevance']} | "
                    f"CR: {row['Context Relevance']} | "
                    f"F: {row['Faithfulness']} | "
                    f"Overall: {row['Overall']}"
                )

            report_lines.append("")
            top_model = df.iloc[0]["Model"]
            report_lines.append(f"Top Model: {top_model}")
            report_lines.append("")
            report_lines.append(f"Selected Answer ({r['model']}):")
            report_lines.append(r["answer"])

        report_text = "\n".join(report_lines)

        st.divider()

        st.download_button(
            label="📄 Download Evaluation Report",
            data=report_text,
            file_name="rag_evaluation_report.txt",
            mime="text/plain"
        )



# -------------------------------
# 📜 View Chat Logs
# -------------------------------
if tabs == "📜 View Chat Logs":
    st.header("📜 View Your Chat Logs")

    log_files = sorted(user_logs_dir.glob("*.txt"))
    if not log_files:
        st.info("No logs found yet. Start chatting first!")
    else:
        selected_log = st.selectbox("Select log file:", [f.name for f in log_files])
        if selected_log:
            with open(user_logs_dir / selected_log, "r", encoding="utf-8") as f:
                st.text(f.read())
